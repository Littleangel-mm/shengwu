import hashlib
import io
from collections.abc import Callable, Sequence
from typing import Any
from uuid import UUID

import pandas as pd
from docx import Document as WordDocument
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from sqlalchemy import delete, func, insert, select, update
from sqlalchemy.orm import Session

from app.core.errors import AppError
from app.db.tables import table
from app.models import ProcessingJob, Project, StoredFile
from app.schemas.report import ReportCreate
from app.schemas.workflow import TaskAccepted
from app.services.prisma import PRISMA_LABELS, PrismaService
from app.services.storage import LocalStorage


class ReportService:
    def __init__(self, db: Session, storage: LocalStorage) -> None:
        self.db = db
        self.storage = storage

    def create(
        self, project_id: UUID, payload: ReportCreate, actor_id: UUID | None
    ) -> TaskAccepted:
        self._validate_source_scope(
            project_id,
            payload.dataset_version_id,
            payload.ml_run_id,
            payload.optimization_run_id,
        )
        reports = table(self.db, "reports")
        version_no = (
            self.db.scalar(
                select(func.max(reports.c.version_no)).where(
                    reports.c.project_id == project_id, reports.c.title == payload.title
                )
            )
            or 0
        ) + 1
        report_id = self.db.execute(
            insert(reports)
            .values(
                project_id=project_id,
                dataset_version_id=payload.dataset_version_id,
                ml_run_id=payload.ml_run_id,
                optimization_run_id=payload.optimization_run_id,
                version_no=version_no,
                title=payload.title,
                status="queued",
                configuration=payload.configuration,
                generated_by=actor_id,
            )
            .returning(reports.c.id)
        ).scalar_one()
        job = ProcessingJob(
            project_id=project_id,
            job_type="generate_report",
            status="queued",
            progress_percent=0,
            current_stage="waiting",
            idempotency_key=f"generate_report:{report_id}",
            requested_config={"report_id": str(report_id)},
            result_summary={},
            requested_by=actor_id,
        )
        self.db.add(job)
        self.db.flush()
        self.db.execute(update(reports).where(reports.c.id == report_id).values(job_id=job.id))
        self.db.commit()
        return TaskAccepted(resource_id=report_id, job_id=job.id)

    def _validate_source_scope(
        self,
        project_id: UUID,
        dataset_version_id: UUID,
        ml_run_id: UUID | None,
        optimization_run_id: UUID | None,
    ) -> None:
        datasets = table(self.db, "datasets")
        versions = table(self.db, "dataset_versions")
        dataset_project_id = self.db.scalar(
            select(datasets.c.project_id)
            .join(versions, versions.c.dataset_id == datasets.c.id)
            .where(
                versions.c.id == dataset_version_id,
                datasets.c.deleted_at.is_(None),
            )
        )
        if dataset_project_id != project_id:
            raise AppError(
                code="report_source_not_found",
                message="报告数据源不存在",
                status_code=404,
            )

        if ml_run_id is not None:
            ml_runs = table(self.db, "ml_runs")
            ml_project_id = self.db.scalar(
                select(ml_runs.c.project_id).where(ml_runs.c.id == ml_run_id)
            )
            if ml_project_id != project_id:
                raise AppError(
                    code="report_source_not_found",
                    message="报告数据源不存在",
                    status_code=404,
                )

        if optimization_run_id is not None:
            optimizations = table(self.db, "optimization_runs")
            optimization_project_id = self.db.scalar(
                select(optimizations.c.project_id).where(optimizations.c.id == optimization_run_id)
            )
            if optimization_project_id != project_id:
                raise AppError(
                    code="report_source_not_found",
                    message="报告数据源不存在",
                    status_code=404,
                )

    def list(self, project_id: UUID) -> list[dict]:
        reports = table(self.db, "reports")
        return [
            dict(row)
            for row in self.db.execute(
                select(reports)
                .where(reports.c.project_id == project_id)
                .order_by(reports.c.created_at.desc())
            ).mappings()
        ]

    def get(self, project_id: UUID, report_id: UUID) -> dict:
        reports = table(self.db, "reports")
        report_assets = table(self.db, "report_assets")
        row = (
            self.db.execute(
                select(reports).where(
                    reports.c.id == report_id,
                    reports.c.project_id == project_id,
                )
            )
            .mappings()
            .one_or_none()
        )
        if not row:
            raise AppError(code="report_not_found", message="报告不存在", status_code=404)
        result = dict(row)
        result["assets"] = [
            dict(asset)
            for asset in self.db.execute(
                select(report_assets)
                .where(report_assets.c.report_id == report_id)
                .order_by(report_assets.c.position, report_assets.c.created_at)
            ).mappings()
        ]
        result["download_ready"] = bool(row["output_file_id"] and row["status"] == "completed")
        return result

    @staticmethod
    def _evidence_truncation_notice() -> str:
        return "证据索引超过 1000 条，本 Word 报告仅展示前 1000 条；完整证据请在系统中查看。"

    @staticmethod
    def _pyplot() -> Any:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False
        return plt

    @classmethod
    def _add_chart(cls, document: Any, figure: Any, title: str) -> None:
        plt = cls._pyplot()

        buffer = io.BytesIO()
        figure.tight_layout()
        figure.savefig(buffer, format="png", dpi=160, bbox_inches="tight")
        plt.close(figure)
        buffer.seek(0)
        document.add_heading(title, level=2)
        document.add_picture(buffer, width=Inches(6.2))

    @staticmethod
    def _apply_chinese_font(document: Any) -> None:
        """确保正文与各级标题使用中文字体，避免 Word 中文回退为方块。"""
        for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
            try:
                style = document.styles[style_name]
            except KeyError:
                continue
            style.font.name = "Microsoft YaHei"
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rfonts.set(qn("w:ascii"), "Microsoft YaHei")
            rfonts.set(qn("w:hAnsi"), "Microsoft YaHei")

    @staticmethod
    def _placeholder(document: Any, text: str) -> None:
        """缺步骤/缺数据时写入中性占位说明，保证章节完整且不崩溃。"""
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.italic = True

    @staticmethod
    def _software_environment() -> Sequence[tuple[str, str]]:
        from importlib import metadata

        packages = [
            ("Python", "runtime"),
            ("scikit-learn", "sklearn"),
            ("pandas", "pandas"),
            ("numpy", "numpy"),
            ("scipy", "scipy"),
            ("matplotlib", "matplotlib"),
            ("python-docx", "docx"),
            ("lightgbm", "lightgbm"),
        ]
        import platform

        rows: list[tuple[str, str]] = []
        for label, dist in packages:
            if dist == "runtime":
                rows.append((label, platform.python_version()))
                continue
            try:
                rows.append((label, metadata.version(dist)))
            except metadata.PackageNotFoundError:
                continue
        return rows

    def _add_flow_diagram(self, document: Any, steps: Sequence[str]) -> None:
        plt = self._pyplot()

        figure, axis = plt.subplots(figsize=(7.2, max(2.0, 0.9 * len(steps))))
        axis.set_xlim(0, 10)
        axis.set_ylim(0, len(steps))
        axis.axis("off")
        centers: list[float] = []
        for index, step in enumerate(reversed(steps)):
            y = index + 0.5
            centers.append(y)
            axis.add_patch(
                plt.Rectangle(
                    (2, y - 0.32),
                    6,
                    0.64,
                    facecolor="#f5efe1",
                    edgecolor="#8a6d3b",
                    linewidth=1.1,
                )
            )
            axis.text(5, y, step, ha="center", va="center", fontsize=10)
        for index in range(len(centers) - 1):
            axis.annotate(
                "",
                xy=(5, centers[index] - 0.32),
                xytext=(5, centers[index + 1] + 0.32),
                arrowprops={"arrowstyle": "-|>", "color": "#8a6d3b", "lw": 1.1},
            )
        self._add_chart(document, figure, "研究流程")

    def generate(self, report_id: UUID, progress: Callable[[float, str], None]) -> dict[str, Any]:
        reports = table(self.db, "reports")
        datasets = table(self.db, "datasets")
        versions = table(self.db, "dataset_versions")
        ml_runs = table(self.db, "ml_runs")
        ml_models = table(self.db, "ml_models")
        ml_metrics = table(self.db, "ml_metrics")
        ml_predictions = table(self.db, "ml_predictions")
        ml_explanations = table(self.db, "ml_explanations")
        optimizations = table(self.db, "optimization_runs")
        candidates = table(self.db, "optimization_candidates")
        dataset_fields = table(self.db, "dataset_fields")
        dataset_rows = table(self.db, "dataset_rows")
        dataset_cells = table(self.db, "dataset_cells")
        cell_evidence = table(self.db, "dataset_cell_evidence")
        conversion_records = table(self.db, "conversion_records")
        documents = table(self.db, "documents")
        pages = table(self.db, "document_pages")
        report_assets = table(self.db, "report_assets")
        report = (
            self.db.execute(select(reports).where(reports.c.id == report_id))
            .mappings()
            .one_or_none()
        )
        if not report:
            raise AppError(code="report_not_found", message="报告任务不存在", status_code=404)
        self._validate_source_scope(
            report["project_id"],
            report["dataset_version_id"],
            report["ml_run_id"],
            report["optimization_run_id"],
        )
        self.db.execute(delete(report_assets).where(report_assets.c.report_id == report_id))
        project = self.db.get(Project, report["project_id"])
        if not project:
            raise AppError(code="project_not_found", message="项目不存在", status_code=404)
        version = (
            self.db.execute(select(versions).where(versions.c.id == report["dataset_version_id"]))
            .mappings()
            .one()
        )
        dataset = (
            self.db.execute(select(datasets).where(datasets.c.id == version["dataset_id"]))
            .mappings()
            .one()
        )
        document = WordDocument()
        styles = document.styles
        styles["Normal"].font.size = Pt(10.5)
        self._apply_chinese_font(document)
        document.add_heading(report["title"], level=0)
        document.add_paragraph(f"项目：{project.name}")
        document.add_paragraph(f"数据集：{dataset['name']}，版本 V{version['version_no']}")
        document.add_paragraph(
            f"数据规模：{version['row_count']} 行，{version['field_count']} 个字段"
        )
        document.add_paragraph(f"数据校验哈希：{version['content_sha256'] or '草稿版本，尚未冻结'}")
        progress(20, "writing_summary_section")

        includes_model = bool(report["ml_run_id"])
        includes_optimization = bool(report["optimization_run_id"])
        document.add_heading("摘要", level=1)
        document.add_paragraph(
            f"本报告基于项目「{project.name}」的冻结数据集 V{version['version_no']}"
            f"（{version['row_count']} 行、{version['field_count']} 个字段）编制，"
            "所有数值均可回溯至原始文献证据。"
            + ("报告包含机器学习建模结果。" if includes_model else "报告未包含机器学习建模结果。")
            + ("并给出目标参数优化推荐与验证方案。" if includes_optimization else "")
        )

        document.add_heading("研究目标", level=1)
        if project.description:
            document.add_paragraph(str(project.description))
        else:
            self._placeholder(
                document, "项目未填写研究目标描述，可在项目设置中补充后重新生成报告。"
            )
        objectives_config = (report["configuration"] or {}).get("objectives")
        if isinstance(objectives_config, list) and objectives_config:
            for objective in objectives_config:
                document.add_paragraph(str(objective), style="List Bullet")

        document.add_heading("研究流程", level=1)
        document.add_paragraph(
            "平台采用统一的可追溯流程：文献上传后自动发现候选字段与关键词，经人工确认后"
            "抽取结构化数据，构建可冻结的数据集版本，训练与评估模型，进行目标参数优化推荐，"
            "并最终生成带证据索引的研究报告。"
        )
        flow_steps = ["文献上传", "字段/关键词自动发现", "人工确认", "结构化抽取", "数据集冻结"]
        if includes_model:
            flow_steps.append("模型训练与评估")
        if includes_optimization:
            flow_steps.append("目标参数优化推荐")
        flow_steps.append("研究报告生成")
        self._add_flow_diagram(document, flow_steps)

        if bool((project.settings or {}).get("enable_prisma")):
            document.add_heading("文献筛选流程（PRISMA 2020）", level=1)
            flow = PrismaService(self.db).get_flow(project.id)
            prisma_data = flow["data"]
            if flow.get("exists"):
                prisma_table = document.add_table(rows=1, cols=2)
                prisma_table.rows[0].cells[0].text = "阶段"
                prisma_table.rows[0].cells[1].text = "数量"
                for key, label in PRISMA_LABELS.items():
                    cells_doc = prisma_table.add_row().cells
                    cells_doc[0].text = label
                    cells_doc[1].text = str(prisma_data.get(key, 0))
                reasons = prisma_data.get("reports_excluded") or []
                if reasons:
                    document.add_paragraph("全文排除原因：")
                    for reason in reasons:
                        document.add_paragraph(
                            f"{reason['reason']}：{reason['count']} 篇", style="List Bullet"
                        )
                if flow.get("notes"):
                    document.add_paragraph(f"检索说明：{flow['notes']}")
                try:
                    import io as _io

                    png = PrismaService.render_diagram(prisma_data)
                    document.add_picture(_io.BytesIO(png), width=Inches(4.6))
                    self.db.execute(
                        insert(report_assets).values(
                            report_id=report_id,
                            asset_type="chart",
                            section_key="prisma_flow",
                            title="PRISMA 流程图",
                            data_payload=prisma_data,
                            position=5,
                        )
                    )
                except Exception:  # noqa: BLE001 - 图渲染失败不应阻断报告
                    self._placeholder(document, "PRISMA 流程图渲染失败，已跳过图示，计数见上表。")
            else:
                self._placeholder(
                    document,
                    "已启用 PRISMA 模块，但尚未录入筛选计数，可在“系统综述”页填写后重新生成。",
                )
        progress(28, "writing_dataset_section")

        document.add_heading("数据处理与可追溯性", level=1)
        document.add_paragraph(
            "系统保留原始值、标准值和建模值三层数据。正式数值均可回溯到原始文献、页码和证据文本。"
        )
        document.add_paragraph("缺失值不进行臆造；图像估计值和人工修改值在数据集中单独标记。")

        source_documents = self.db.execute(
            select(documents.c.title, documents.c.publication_year, documents.c.publication_name)
            .where(documents.c.project_id == project.id, documents.c.deleted_at.is_(None))
            .order_by(documents.c.created_at)
        ).all()
        document.add_heading("文献来源与解析范围", level=1)
        document.add_paragraph(f"本项目共纳入 {len(source_documents)} 篇文献。")
        for title, year, publication in source_documents[:200]:
            document.add_paragraph(
                f"{title or '未命名文献'}；{year or '年份未知'}；{publication or '来源未知'}",
                style="List Bullet",
            )

        field_rows = (
            self.db.execute(
                select(dataset_fields)
                .where(dataset_fields.c.dataset_version_id == version["id"])
                .order_by(dataset_fields.c.position)
            )
            .mappings()
            .all()
        )
        document.add_heading("字段、抽取与单位规则", level=1)
        for field in field_rows:
            document.add_paragraph(
                f"{field['display_name']} ({field['field_key']})；类型={field['data_type']}；"
                f"角色={field['semantic_role']}；校验={field['validation_rules']}",
                style="List Bullet",
            )
        row_ids = list(
            self.db.scalars(
                select(dataset_rows.c.id).where(
                    dataset_rows.c.dataset_version_id == version["id"],
                    dataset_rows.c.is_deleted.is_(False),
                )
            ).all()
        )
        numeric_fields = {
            row["id"]: row["display_name"]
            for row in field_rows
            if row["data_type"] in {"number", "integer", "float", "decimal", "range"}
        }
        matrix_values: dict[UUID, dict[str, float]] = {}
        if row_ids and numeric_fields:
            for row_id, field_id, value in self.db.execute(
                select(
                    dataset_cells.c.row_id, dataset_cells.c.field_id, dataset_cells.c.value_number
                ).where(
                    dataset_cells.c.row_id.in_(row_ids),
                    dataset_cells.c.field_id.in_(list(numeric_fields)),
                    dataset_cells.c.value_number.is_not(None),
                    dataset_cells.c.is_missing.is_(False),
                )
            ):
                matrix_values.setdefault(row_id, {})[numeric_fields[field_id]] = float(value)
        numeric_frame = pd.DataFrame(matrix_values.values())
        if not numeric_frame.empty:
            plt = self._pyplot()

            columns = list(numeric_frame.columns[:6])
            figure, axes = plt.subplots(len(columns), 1, figsize=(7, max(2.4, 2.1 * len(columns))))
            axes_list = [axes] if len(columns) == 1 else list(axes)
            for axis, column in zip(axes_list, columns, strict=False):
                axis.hist(
                    numeric_frame[column].dropna(), bins=min(20, max(5, len(numeric_frame) // 2))
                )
                axis.set_title(column)
                axis.set_ylabel("Count")
            self._add_chart(document, figure, "数据分布")
            self.db.execute(
                insert(report_assets).values(
                    report_id=report_id,
                    asset_type="chart",
                    section_key="data_distribution",
                    title="数据分布",
                    data_payload={"fields": columns},
                    position=10,
                )
            )
            if len(numeric_frame.columns) >= 2:
                correlation = numeric_frame.corr(numeric_only=True)
                figure, axis = plt.subplots(figsize=(7, 5.5))
                image = axis.imshow(correlation, vmin=-1, vmax=1, cmap="coolwarm")
                axis.set_xticks(
                    range(len(correlation.columns)), correlation.columns, rotation=45, ha="right"
                )
                axis.set_yticks(range(len(correlation.index)), correlation.index)
                figure.colorbar(image, ax=axis, label="Correlation")
                self._add_chart(document, figure, "相关性矩阵")
                self.db.execute(
                    insert(report_assets).values(
                        report_id=report_id,
                        asset_type="chart",
                        section_key="correlation",
                        title="相关性矩阵",
                        data_payload={"matrix": correlation.fillna(0).to_dict()},
                        position=20,
                    )
                )

        if report["ml_run_id"]:
            run = (
                self.db.execute(select(ml_runs).where(ml_runs.c.id == report["ml_run_id"]))
                .mappings()
                .one_or_none()
            )
            if run:
                document.add_heading("机器学习模型", level=1)
                model_rows = (
                    self.db.execute(
                        select(ml_models)
                        .where(ml_models.c.ml_run_id == run["id"])
                        .order_by(ml_models.c.model_no)
                    )
                    .mappings()
                    .all()
                )
                document.add_heading("建模设计", level=2)
                document.add_paragraph(
                    f"任务类型：{run['task_type']}；切分策略：{run['split_strategy']}；随机种子：{run['random_seed']}"
                )
                algorithms = sorted(
                    {
                        str(model.get("algorithm") or model.get("display_name"))
                        for model in model_rows
                    }
                )
                if algorithms:
                    document.add_paragraph("候选算法：" + "、".join(algorithms))
                document.add_paragraph(
                    "为避免数据泄漏，特征标准化仅在训练折内拟合并应用于验证/测试；"
                    "多目标建模时对各目标做训练集内归一化后加权合成；"
                    "交叉验证按可用分组信息自动选择策略，退化时在训练摘要中标注。"
                )
                document.add_heading("模型对比", level=2)
                table_doc = document.add_table(rows=1, cols=4)
                for index, title in enumerate(["模型", "状态", "是否选中", "测试指标"]):
                    table_doc.rows[0].cells[index].text = title
                for model in model_rows:
                    metric_rows = (
                        self.db.execute(
                            select(ml_metrics).where(
                                ml_metrics.c.ml_model_id == model["id"],
                                ml_metrics.c.split_name == "test",
                            )
                        )
                        .mappings()
                        .all()
                    )
                    row = table_doc.add_row().cells
                    row[0].text = model["display_name"]
                    row[1].text = model["status"]
                    row[2].text = "是" if model["is_selected"] else "否"
                    row[3].text = ", ".join(
                        f"{item['metric_name']}={item['metric_value']:.4f}"
                        for item in metric_rows
                        if item["metric_value"] is not None
                    )
                    if model["status"] != "completed" or not model["is_selected"]:
                        # 仅展开“用户选中且训练成功”的模型明细，失败/未选模型只在对比表体现。
                        continue
                    explanation_rows = (
                        self.db.execute(
                            select(ml_explanations).where(
                                ml_explanations.c.ml_model_id == model["id"]
                            )
                        )
                        .mappings()
                        .all()
                    )
                    for explanation in explanation_rows:
                        if explanation["method"] not in {"shap", "permutation_importance"}:
                            continue
                        document.add_heading(
                            f"{model['display_name']} - {explanation['method']}", level=2
                        )
                        features = (explanation["explanation_data"] or {}).get("features", [])
                        for feature in sorted(
                            features,
                            key=lambda item: abs(
                                item.get("mean_absolute_shap", item.get("importance_mean", 0))
                            ),
                            reverse=True,
                        )[:20]:
                            score = feature.get(
                                "mean_absolute_shap", feature.get("importance_mean", 0)
                            )
                            document.add_paragraph(f"{feature.get('name')}: {score:.6g}")
                completed_models = [m for m in model_rows if m["status"] == "completed"]
                selected_model = next(
                    (m for m in completed_models if m["is_selected"]),
                    completed_models[0] if completed_models else None,
                )
                if not selected_model:
                    self._placeholder(
                        document,
                        "本次运行没有训练成功的模型，故未生成模型图表；请检查数据规模或算法配置后重试。",
                    )
                if selected_model:
                    document.add_heading("模型结果与公式", level=2)
                    document.add_paragraph(
                        f"入选模型：{selected_model['display_name']}"
                        f"（算法：{selected_model.get('algorithm') or '未标注'}）。"
                    )
                    selected_metrics = (
                        self.db.execute(
                            select(ml_metrics).where(
                                ml_metrics.c.ml_model_id == selected_model["id"],
                                ml_metrics.c.split_name == "test",
                            )
                        )
                        .mappings()
                        .all()
                    )
                    if selected_metrics:
                        document.add_paragraph(
                            "测试集指标："
                            + "，".join(
                                f"{item['metric_name']}={item['metric_value']:.4f}"
                                for item in selected_metrics
                                if item["metric_value"] is not None
                            )
                        )
                    else:
                        self._placeholder(document, "入选模型暂无测试集指标记录。")
                    document.add_paragraph(
                        "模型采用 scikit-learn 标准实现，其解析形式由所选算法决定"
                        "（如线性/岭回归为特征线性加权，树/集成模型为分段常数的加权组合），"
                        "特征贡献以特征重要性图与上文数值列表呈现，不以人工设定的经验公式替代。"
                    )
                    prediction_rows = self.db.execute(
                        select(
                            ml_predictions.c.actual_value,
                            ml_predictions.c.predicted_value,
                            ml_predictions.c.residual_value,
                        ).where(
                            ml_predictions.c.ml_model_id == selected_model["id"],
                            ml_predictions.c.split_name == "test",
                        )
                    ).all()
                    actual = [
                        float(row.actual_value["value"])
                        for row in prediction_rows
                        if row.actual_value and row.actual_value.get("value") is not None
                    ]
                    predicted = [
                        float(row.predicted_value["value"])
                        for row in prediction_rows
                        if row.predicted_value and row.predicted_value.get("value") is not None
                    ]
                    residuals = [
                        float(row.residual_value)
                        for row in prediction_rows
                        if row.residual_value is not None
                    ]
                    if actual and len(actual) == len(predicted):
                        plt = self._pyplot()

                        figure, axes = plt.subplots(1, 2, figsize=(9, 4))
                        axes[0].scatter(actual, predicted, alpha=0.8)
                        lower = min([*actual, *predicted])
                        upper = max([*actual, *predicted])
                        axes[0].plot([lower, upper], [lower, upper], linestyle="--", color="black")
                        axes[0].set_xlabel("Actual")
                        axes[0].set_ylabel("Predicted")
                        axes[0].set_title("Actual vs Predicted")
                        if residuals:
                            axes[1].scatter(predicted[: len(residuals)], residuals, alpha=0.8)
                            axes[1].axhline(0, linestyle="--", color="black")
                            axes[1].set_xlabel("Predicted")
                            axes[1].set_ylabel("Residual")
                            axes[1].set_title("Residuals")
                        self._add_chart(document, figure, "实测、预测与残差")
                        self.db.execute(
                            insert(report_assets).values(
                                report_id=report_id,
                                asset_type="chart",
                                section_key="prediction_residual",
                                title="实测、预测与残差",
                                data_payload={"model_id": str(selected_model["id"])},
                                position=30,
                            )
                        )
                    selected_explanation = (
                        self.db.execute(
                            select(ml_explanations).where(
                                ml_explanations.c.ml_model_id == selected_model["id"],
                                ml_explanations.c.method.in_(["shap", "permutation_importance"]),
                            )
                        )
                        .mappings()
                        .first()
                    )
                    if selected_explanation:
                        features = (selected_explanation["explanation_data"] or {}).get(
                            "features", []
                        )
                        ranked = sorted(
                            features,
                            key=lambda item: abs(
                                item.get("mean_absolute_shap", item.get("importance_mean", 0))
                            ),
                            reverse=True,
                        )[:20]
                        if ranked:
                            plt = self._pyplot()

                            names = [str(item.get("name")) for item in reversed(ranked)]
                            scores = [
                                float(
                                    item.get("mean_absolute_shap", item.get("importance_mean", 0))
                                )
                                for item in reversed(ranked)
                            ]
                            figure, axis = plt.subplots(figsize=(7, max(3, len(names) * 0.3)))
                            axis.barh(names, scores)
                            axis.set_xlabel(selected_explanation["method"])
                            self._add_chart(document, figure, "特征重要性")
                            self.db.execute(
                                insert(report_assets).values(
                                    report_id=report_id,
                                    asset_type="chart",
                                    section_key="feature_importance",
                                    title="特征重要性",
                                    data_payload={
                                        "model_id": str(selected_model["id"]),
                                        "features": ranked,
                                    },
                                    position=40,
                                )
                            )
        progress(55, "writing_model_section")

        if report["optimization_run_id"]:
            optimization = (
                self.db.execute(
                    select(optimizations).where(optimizations.c.id == report["optimization_run_id"])
                )
                .mappings()
                .one_or_none()
            )
            if optimization:
                document.add_heading("目标参数优化候选", level=1)
                document.add_paragraph(
                    f"方法：{optimization['method']}；目标：{optimization['objective_config']}"
                )
                rows = (
                    self.db.execute(
                        select(candidates)
                        .where(candidates.c.optimization_run_id == optimization["id"])
                        .order_by(candidates.c.rank_no)
                        .limit(20)
                    )
                    .mappings()
                    .all()
                )
                for item in rows:
                    document.add_paragraph(
                        f"#{item['rank_no']} 输入={item['input_values']}，预测={item['predicted_values']}，目标分数={item['objective_score']}"
                    )
                summary = optimization["result_summary"] or {}
                document.add_heading("验证实验方案", level=1)
                validation_plan = summary.get("validation_plan") or {}
                control = validation_plan.get("control")
                recommendations = validation_plan.get("recommendations") or []
                if control or recommendations:
                    if control:
                        document.add_paragraph(f"对照组基线：{control}")
                    document.add_paragraph("建议验证的推荐条件（按排名）：")
                    for index, group in enumerate(recommendations[:10], start=1):
                        document.add_paragraph(f"实验 {index}：{group}", style="List Bullet")
                    document.add_paragraph(
                        "建议每个条件设置重复实验，并在真实环境复测；超出训练数据适用域的候选需谨慎解读。"
                    )
                else:
                    self._placeholder(
                        document,
                        "本次优化未生成结构化验证方案，请依据上表推荐条件设计对照与重复实验。",
                    )
        document.add_heading("单位转换记录", level=1)
        conversion_rows = self.db.execute(
            select(
                conversion_records.c.source_value,
                conversion_records.c.target_value,
                conversion_records.c.source_unit_text,
                conversion_records.c.formula_used,
                conversion_records.c.status,
            )
            .join(dataset_cells, dataset_cells.c.id == conversion_records.c.dataset_cell_id)
            .join(dataset_rows, dataset_rows.c.id == dataset_cells.c.row_id)
            .where(dataset_rows.c.dataset_version_id == version["id"])
            .limit(500)
        ).all()
        if conversion_rows:
            conversion_table = document.add_table(rows=1, cols=5)
            for index, title in enumerate(["原值", "目标值", "原单位", "公式", "状态"]):
                conversion_table.rows[0].cells[index].text = title
            for conversion_item in conversion_rows:
                cells_doc = conversion_table.add_row().cells
                cells_doc[0].text = str(conversion_item.source_value)
                cells_doc[1].text = str(conversion_item.target_value)
                cells_doc[2].text = str(conversion_item.source_unit_text or "")
                cells_doc[3].text = conversion_item.formula_used
                cells_doc[4].text = conversion_item.status
        else:
            document.add_paragraph("本版本无单位转换记录。")

        document.add_heading("证据索引", level=1)
        evidence_rows = self.db.execute(
            select(
                documents.c.title,
                pages.c.page_no,
                cell_evidence.c.evidence_text,
                cell_evidence.c.bbox,
            )
            .join(dataset_cells, dataset_cells.c.id == cell_evidence.c.dataset_cell_id)
            .join(dataset_rows, dataset_rows.c.id == dataset_cells.c.row_id)
            .join(versions, versions.c.id == dataset_rows.c.dataset_version_id)
            .join(documents, documents.c.id == dataset_rows.c.source_document_id)
            .join(pages, pages.c.id == cell_evidence.c.page_id)
            .where(dataset_rows.c.dataset_version_id == version["id"])
            .order_by(documents.c.title, pages.c.page_no)
            .limit(1001)
        ).all()
        evidence_truncated = len(evidence_rows) > 1000
        evidence_rows = evidence_rows[:1000]
        if evidence_truncated:
            document.add_paragraph(self._evidence_truncation_notice())
        if evidence_rows:
            evidence_table = document.add_table(rows=1, cols=4)
            for index, title in enumerate(["文献", "页码", "原文证据", "坐标"]):
                evidence_table.rows[0].cells[index].text = title
            for evidence_item in evidence_rows:
                cells_doc = evidence_table.add_row().cells
                cells_doc[0].text = evidence_item.title or ""
                cells_doc[1].text = str(evidence_item.page_no)
                cells_doc[2].text = evidence_item.evidence_text
                cells_doc[3].text = str(evidence_item.bbox or "")
        else:
            document.add_paragraph("本版本未关联证据索引。")
        document.add_heading("局限性", level=1)
        document.add_paragraph(
            "模型输出是基于已审核文献数据的统计预测，不替代真实实验和专业判断。超出训练数据范围的候选条件应重新验证。"
        )

        document.add_heading("可复现性", level=1)
        document.add_paragraph(
            f"数据集版本：V{version['version_no']}；"
            f"内容校验哈希：{version['content_sha256'] or '草稿版本，尚未冻结'}。"
        )
        if report["ml_run_id"]:
            document.add_paragraph(
                "建模随机种子、切分与交叉验证策略见“建模设计”章节；"
                "在相同数据版本与随机种子下可复现建模结果。"
            )
        document.add_paragraph(
            "所有正式数值均保留原始值、标准值、建模值三层，并关联文献页码与证据文本，"
            "可据此复核与复现数据来源。"
        )

        document.add_heading("开源软件与方法", level=1)
        software_rows = self._software_environment()
        if software_rows:
            software_table = document.add_table(rows=1, cols=2)
            software_table.rows[0].cells[0].text = "软件/库"
            software_table.rows[0].cells[1].text = "版本"
            for name, version_text in software_rows:
                cells_doc = software_table.add_row().cells
                cells_doc[0].text = name
                cells_doc[1].text = version_text
        else:
            self._placeholder(document, "未能获取软件版本信息。")

        document.add_heading("发表建议", level=1)
        document.add_paragraph(
            "撰写论文时，建议在方法部分引用上述软件与版本，并按 PRISMA 报告文献筛选流程；"
            "在结果部分呈现模型评估指标、特征重要性与优化推荐，并附本报告的数据版本哈希以支持复现；"
            "在讨论部分说明适用域与局限性，对关键推荐条件给出实验验证计划。"
        )
        progress(80, "saving_report")

        buffer = io.BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        saved = self.storage.save_bytes(
            project.id,
            category="reports",
            extension="docx",
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        stored = StoredFile(
            organization_id=project.organization_id,
            project_id=project.id,
            storage_provider="local",
            storage_key=saved.storage_key,
            original_name=f"{report['title']}-v{report['version_no']}.docx",
            safe_name=saved.safe_name,
            extension="docx",
            media_type=saved.media_type,
            byte_size=saved.byte_size,
            sha256=saved.sha256,
            purpose="report",
            security_status="generated",
            metadata_json={},
        )
        self.db.add(stored)
        self.db.flush()
        self.db.execute(
            update(reports)
            .where(reports.c.id == report_id)
            .values(
                status="completed",
                output_file_id=stored.id,
                content_sha256=hashlib.sha256(content).hexdigest(),
                completed_at=func.now(),
            )
        )
        self.db.commit()
        return {"report_id": str(report_id), "file_id": str(stored.id), "sha256": saved.sha256}

    def output_path(self, project_id: UUID, report_id: UUID):
        reports = table(self.db, "reports")
        row = (
            self.db.execute(
                select(reports).where(reports.c.id == report_id, reports.c.project_id == project_id)
            )
            .mappings()
            .one_or_none()
        )
        if not row or not row["output_file_id"]:
            raise AppError(code="report_not_ready", message="报告尚未生成", status_code=409)
        stored = self.db.get(StoredFile, row["output_file_id"])
        if not stored:
            raise AppError(code="report_file_not_found", message="报告文件不存在", status_code=404)
        return self.storage.path_for_key(stored.storage_key), stored.original_name
