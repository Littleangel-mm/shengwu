import io
import os
from uuid import uuid4

import pytest
from docx import Document as WordDocument
from sqlalchemy import delete, insert, update

pytestmark = pytest.mark.skipif(
    os.getenv("RUN_INTEGRATION_TESTS") != "1",
    reason="requires an isolated PostgreSQL acceptance database",
)


def _seed_minimal_report(db, *, enable_prisma: bool, with_prisma_counts: bool):
    from app.db.tables import table

    organizations = table(db, "organizations")
    projects = table(db, "projects")
    datasets = table(db, "datasets")
    versions = table(db, "dataset_versions")
    fields = table(db, "dataset_fields")
    reports = table(db, "reports")

    suffix = uuid4().hex[:10]
    org_id = db.execute(
        insert(organizations)
        .values(name=f"E2E Org {suffix}", slug=f"e2e-org-{suffix}")
        .returning(organizations.c.id)
    ).scalar_one()
    project_id = db.execute(
        insert(projects)
        .values(
            organization_id=org_id,
            name="E2E Report Project",
            slug=f"e2e-report-{suffix}",
            description="验证报告章节完整性的端到端项目。",
            default_language="zh-CN",
            status="active",
            settings={"enable_prisma": enable_prisma},
        )
        .returning(projects.c.id)
    ).scalar_one()
    dataset_id = db.execute(
        insert(datasets)
        .values(project_id=project_id, name="E2E 数据集")
        .returning(datasets.c.id)
    ).scalar_one()
    version_id = db.execute(
        insert(versions)
        .values(
            dataset_id=dataset_id,
            version_no=1,
            status="draft",
            row_count=0,
            field_count=2,
        )
        .returning(versions.c.id)
    ).scalar_one()
    db.execute(
        insert(fields).values(
            dataset_version_id=version_id,
            field_key="temperature",
            display_name="温度",
            data_type="number",
            semantic_role="feature",
            position=0,
        )
    )
    db.execute(
        insert(fields).values(
            dataset_version_id=version_id,
            field_key="yield_rate",
            display_name="产率",
            data_type="number",
            semantic_role="target",
            position=1,
        )
    )
    db.execute(
        update(versions)
        .where(versions.c.id == version_id)
        .values(status="frozen", content_sha256="0" * 64)
    )
    if enable_prisma and with_prisma_counts:
        prisma = table(db, "prisma_flows")
        db.execute(
            insert(prisma).values(
                project_id=project_id,
                data={
                    "identified_databases": 120,
                    "identified_registers": 0,
                    "duplicates_removed": 20,
                    "records_screened": 100,
                    "records_excluded": 60,
                    "reports_sought": 40,
                    "reports_not_retrieved": 2,
                    "reports_assessed": 38,
                    "studies_included": 30,
                    "reports_excluded": [{"reason": "无对照组", "count": 8}],
                },
                notes="检索时间窗口 2015-2025。",
            )
        )
    report_id = db.execute(
        insert(reports)
        .values(
            project_id=project_id,
            dataset_version_id=version_id,
            version_no=1,
            title="端到端测试报告",
            status="queued",
            configuration={},
        )
        .returning(reports.c.id)
    ).scalar_one()
    db.commit()
    return org_id, project_id, report_id


def _render(enable_prisma: bool, with_prisma_counts: bool) -> list[str]:
    from app.core.config import get_settings
    from app.db.session import SessionLocal
    from app.db.tables import table
    from app.services.report import ReportService
    from app.services.storage import LocalStorage

    db = SessionLocal()
    org_id = None
    try:
        org_id, project_id, report_id = _seed_minimal_report(
            db, enable_prisma=enable_prisma, with_prisma_counts=with_prisma_counts
        )
        service = ReportService(db, LocalStorage(get_settings()))
        result = service.generate(report_id, lambda *_: None)
        assert result["sha256"]
        path, _ = service.output_path(project_id, report_id)
        with open(path, "rb") as handle:
            document = WordDocument(io.BytesIO(handle.read()))
        return [p.text for p in document.paragraphs]
    finally:
        if org_id is not None:
            organizations = table(db, "organizations")
            db.execute(delete(organizations).where(organizations.c.id == org_id))
            db.commit()
        db.close()


def test_report_contains_all_required_sections_without_model() -> None:
    texts = _render(enable_prisma=False, with_prisma_counts=False)
    joined = "\n".join(texts)
    for heading in (
        "摘要",
        "研究目标",
        "研究流程",
        "数据处理与可追溯性",
        "字段、抽取与单位规则",
        "局限性",
        "可复现性",
        "开源软件与方法",
        "发表建议",
    ):
        assert heading in joined, f"missing section: {heading}"


def test_report_embeds_prisma_when_enabled() -> None:
    texts = _render(enable_prisma=True, with_prisma_counts=True)
    joined = "\n".join(texts)
    assert "PRISMA" in joined
    assert "检索时间窗口" in joined


def test_report_prisma_placeholder_when_counts_missing() -> None:
    texts = _render(enable_prisma=True, with_prisma_counts=False)
    joined = "\n".join(texts)
    assert "尚未录入筛选计数" in joined


def test_csv_document_parses_into_table_cells() -> None:
    import hashlib

    from sqlalchemy import delete as sa_delete
    from sqlalchemy import select

    from app.core.config import get_settings
    from app.db.session import SessionLocal
    from app.db.tables import table
    from app.models import (
        Document,
        DocumentVersion,
        Organization,
        Project,
        StoredFile,
    )
    from app.services.parser import DocumentParser
    from app.services.storage import LocalStorage

    db = SessionLocal()
    org = None
    try:
        settings = get_settings()
        storage = LocalStorage(settings)
        storage.ensure_root()
        suffix = uuid4().hex[:10]
        org = Organization(name=f"CSV Org {suffix}", slug=f"csv-org-{suffix}")
        db.add(org)
        db.flush()
        project = Project(
            organization_id=org.id, name="CSV Project", slug=f"csv-proj-{suffix}"
        )
        db.add(project)
        db.flush()
        content = b"temperature,yield\n30,0.82\n45,0.91\n"
        key = f"uploads/test/{suffix}.csv"
        destination = storage.path_for_key(key)
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(content)
        stored = StoredFile(
            organization_id=org.id,
            project_id=project.id,
            storage_provider="local",
            storage_key=key,
            original_name="data.csv",
            safe_name=f"{suffix}.csv",
            extension="csv",
            media_type="text/csv",
            byte_size=len(content),
            sha256=hashlib.sha256(content).hexdigest(),
            purpose="upload",
            security_status="clean",
        )
        db.add(stored)
        db.flush()
        document = Document(project_id=project.id, title="CSV 数据")
        db.add(document)
        db.flush()
        version = DocumentVersion(
            document_id=document.id, version_no=1, source_file_id=stored.id
        )
        db.add(version)
        db.commit()

        summary = DocumentParser(db, storage).parse(version.id, lambda *_: None)
        assert summary["tables"] == 1
        assert summary["cells"] == 6

        tables_table = table(db, "document_tables")
        cells_table = table(db, "document_table_cells")
        table_id = db.scalar(
            select(tables_table.c.id).where(
                tables_table.c.document_version_id == version.id
            )
        )
        header = db.scalar(
            select(cells_table.c.raw_text).where(
                cells_table.c.table_id == table_id,
                cells_table.c.row_index == 0,
                cells_table.c.column_index == 0,
            )
        )
        assert header == "temperature"
    finally:
        if org is not None:
            # 先删文献（级联清理版本），再删组织，避开 source_file_id 的 RESTRICT。
            documents = table(db, "documents")
            organizations = table(db, "organizations")
            db.execute(sa_delete(documents).where(documents.c.project_id == project.id))
            db.execute(sa_delete(organizations).where(organizations.c.id == org.id))
            db.commit()
        db.close()


def test_report_lineage_aggregates_chain() -> None:
    from app.db.session import SessionLocal
    from app.db.tables import table
    from app.services.lineage import LineageService

    db = SessionLocal()
    org_id = None
    try:
        org_id, project_id, report_id = _seed_minimal_report(
            db, enable_prisma=False, with_prisma_counts=False
        )
        lineage = LineageService(db).report_lineage(project_id, report_id)
        assert lineage["report"]["title"] == "端到端测试报告"
        assert lineage["dataset_version"]["content_sha256"] == "0" * 64
        assert lineage["ml_run"] is None
        assert lineage["ml_models"] == []
        assert isinstance(lineage["hash_chain"], list)
        stages = {item["stage"] for item in lineage["hash_chain"]}
        assert "冻结数据集" in stages
    finally:
        if org_id is not None:
            organizations = table(db, "organizations")
            from sqlalchemy import delete

            db.execute(delete(organizations).where(organizations.c.id == org_id))
            db.commit()
        db.close()
