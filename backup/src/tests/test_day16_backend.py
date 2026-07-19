from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from app.main import app
from app.services.extraction import convert_parsed_value, normalize_unit_token
from app.services.parser import DocumentParser


def test_docx_body_order_and_embedded_image_relationships() -> None:
    document = Document()
    document.add_paragraph("before")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "inside"
    paragraph = document.add_paragraph("after")
    image = BytesIO()
    Image.new("RGB", (2, 2), "white").save(image, format="PNG")
    image.seek(0)
    paragraph.add_run().add_picture(image)

    items = DocumentParser._docx_body_items(document)

    assert [kind for kind, _ in items[:3]] == ["paragraph", "table", "paragraph"]
    assert items[0][1].text == "before"
    assert items[1][1].cell(0, 0).text == "inside"
    assert items[2][1].text == "after"
    assert len(DocumentParser._docx_image_relationship_ids(items[2][1])) == 1


def test_xlsx_preserves_merged_ranges_formula_and_cached_value() -> None:
    formula_workbook = Workbook()
    formula_sheet = formula_workbook.active
    formula_sheet.merge_cells("A1:B1")
    formula_sheet["A1"] = "header"
    formula_sheet["A2"] = 1
    formula_sheet["B2"] = "=A2+2"
    value_workbook = Workbook()
    value_sheet = value_workbook.active
    value_sheet["A1"] = "header"
    value_sheet["A2"] = 1
    value_sheet["B2"] = 3

    rows, formulas, merged_ranges = DocumentParser._xlsx_sheet_data(
        formula_sheet, value_sheet
    )

    assert merged_ranges == ["A1:B1"]
    assert rows[1] == ["1", "3"]
    assert formulas == [{"coordinate": "B2", "formula": "=A2+2", "cached_value": 3}]


def test_complex_pdf_table_structure_tracks_repeated_merged_geometry() -> None:
    merged_header = (0.0, 0.0, 200.0, 20.0)
    found = SimpleNamespace(
        rows=[
            SimpleNamespace(cells=[merged_header, merged_header]),
            SimpleNamespace(
                cells=[(0.0, 20.0, 100.0, 40.0), (100.0, 20.0, 200.0, 40.0)]
            ),
            SimpleNamespace(
                cells=[(0.0, 40.0, 100.0, 60.0), (100.0, 40.0, 200.0, 60.0)]
            ),
        ],
        header=SimpleNamespace(names=["A", "B"], external=False),
    )
    extracted = [["Measurement", None], ["Yield (g/L)", "Temperature (℃)"], ["4.2", "30"]]

    structure = DocumentParser._table_structure(found, extracted)

    assert structure["header_rows"] == 2
    assert structure["merged_cells"] == [
        {"row": 0, "column": 0, "row_span": 1, "column_span": 2}
    ]
    assert structure["units"] == {"0": "g/L", "1": "℃"}


@pytest.mark.parametrize(
    ("parsed", "expected"),
    [
        ({"type": "number", "value": 25.0}, {"type": "number", "value": 298.15}),
        (
            {"type": "range", "min": 0.0, "max": 10.0, "mid": 5.0},
            {"type": "range", "min": 32.0, "max": 50.0, "mid": 41.0},
        ),
        (
            {"type": "mean_sd", "mean": 10.0, "sd": 2.0},
            {"type": "mean_sd", "mean": 50.0, "sd": 3.6},
        ),
    ],
)
def test_conversion_applies_offset_without_shifting_standard_deviation(
    parsed: dict, expected: dict
) -> None:
    multiplier, offset = ((1.0, 273.15) if parsed["type"] == "number" else (1.8, 32.0))
    assert convert_parsed_value(parsed, multiplier, offset) == pytest.approx(expected)


def test_unit_token_normalization_matches_temperature_aliases() -> None:
    assert normalize_unit_token(" ℃ ") == normalize_unit_token("°C")


def test_conversion_confirmation_routes_are_exposed() -> None:
    paths = app.openapi()["paths"]
    base = "/api/v1/projects/{project_id}/extraction-runs/{run_id}/records/{record_id}"
    assert f"{base}/conversions" in paths
    assert f"{base}/conversions/{{conversion_id}}/confirm" in paths
