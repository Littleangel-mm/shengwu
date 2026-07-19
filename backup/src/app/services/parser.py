import mimetypes
import re
import tempfile
from collections.abc import Callable
from pathlib import Path
from typing import Any
from uuid import UUID

import fitz
from docx import Document as DocxDocument
from openpyxl import load_workbook
from sqlalchemy import delete, insert, select, update
from sqlalchemy.orm import Session

from app.core.errors import AppError
from app.db.tables import table
from app.models import Document, DocumentVersion, Project, StoredFile
from app.services.ocr import OCRImage, OCRPage, PaddleOCRService
from app.services.storage import LocalStorage

ProgressCallback = Callable[[float, str], None]


class DocumentParser:
    def __init__(self, db: Session, storage: LocalStorage) -> None:
        self.db = db
        self.storage = storage

    @staticmethod
    def _bbox(value: Any) -> list[float] | None:
        if value is None:
            return None
        try:
            return [round(float(item), 3) for item in value]
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _nearby_text(
        blocks: list[Any], bbox: Any, *, above: bool, distance: float = 90
    ) -> str | None:
        x0, y0, x1, y1 = [float(item) for item in bbox]
        candidates = []
        for block in blocks:
            if len(block) < 5 or not str(block[4]).strip():
                continue
            bx0, by0, bx1, by1 = [float(item) for item in block[:4]]
            horizontal_overlap = max(0.0, min(x1, bx1) - max(x0, bx0))
            if horizontal_overlap <= 0:
                continue
            gap = y0 - by1 if above else by0 - y1
            if 0 <= gap <= distance:
                candidates.append((gap, str(block[4]).strip()))
        return min(candidates, default=(0, None), key=lambda item: item[0])[1]

    @staticmethod
    def _table_structure(found: Any, extracted: list[list[Any]]) -> dict[str, Any]:
        header_rows = 1
        for index, row in enumerate(extracted[: min(5, len(extracted))]):
            values = [str(value or "").strip() for value in row]
            numeric = sum(bool(re.search(r"\d", value)) for value in values if value)
            nonempty = sum(bool(value) for value in values)
            if index > 0 and nonempty and numeric / nonempty >= 0.6:
                break
            header_rows = index + 1
        column_count = max((len(row) for row in extracted), default=0)
        header_paths = []
        for column in range(column_count):
            path = []
            last = None
            for row_index in range(header_rows):
                row = extracted[row_index] if row_index < len(extracted) else []
                value = str(row[column] or "").strip() if column < len(row) else ""
                if value:
                    last = value
                if last and last not in path:
                    path.append(last)
            header_paths.append(path)
        merged: list[dict[str, int]] = []
        for row_index, row in enumerate(getattr(found, "rows", [])):
            cells = list(getattr(row, "cells", []))
            for column_index, cell in enumerate(cells):
                if cell is None:
                    continue
                column_span = 1
                while (
                    column_index + column_span < len(cells)
                    and cells[column_index + column_span] is None
                ):
                    column_span += 1
                row_span = 1
                while row_index + row_span < len(found.rows):
                    below = list(found.rows[row_index + row_span].cells)
                    if column_index >= len(below) or below[column_index] is not None:
                        break
                    row_span += 1
                if row_span > 1 or column_span > 1:
                    merged.append(
                        {
                            "row": row_index,
                            "column": column_index,
                            "row_span": row_span,
                            "column_span": column_span,
                        }
                    )
        units = {}
        for column, path in enumerate(header_paths):
            match = re.search(
                r"(?:\(|（|/)(%|℃|°C|K|d|h|min|s|mg/L|g/L|mg/g|g/kg|ppm|pH)(?:\)|）)?",
                " / ".join(path),
                re.I,
            )
            if match:
                units[str(column)] = match.group(1)
        return {
            "rows": extracted,
            "header_rows": header_rows,
            "header_names": list(getattr(found.header, "names", []) or []),
            "header_external": bool(getattr(found.header, "external", False)),
            "header_paths": header_paths,
            "merged_cells": merged,
            "units": units,
        }

    @staticmethod
    def _figure_metadata(lines: list[Any], width: float, height: float) -> dict[str, Any]:
        labels = [
            {"text": line.text, "confidence": line.confidence, "bbox": line.bbox} for line in lines
        ]
        numeric_pattern = re.compile(
            r"^[-+]?\d+(?:\.\d+)?(?:\s*(?:±|\+/-)\s*\d+(?:\.\d+)?)?\s*(?:%|℃|°C|K|d|h|min|s|mg/L|g/L|ppm|pH)?$",
            re.I,
        )
        direct_values = [
            item
            for item in labels
            if numeric_pattern.fullmatch(item["text"].strip())
            and item["bbox"][0] > width * 0.12
            and item["bbox"][1] < height * 0.82
        ]
        x_labels = [
            item
            for item in labels
            if item["bbox"][1] >= height * 0.72
            and not numeric_pattern.fullmatch(item["text"].strip())
        ]
        y_labels = [
            item
            for item in labels
            if item["bbox"][0] <= width * 0.22
            and not numeric_pattern.fullmatch(item["text"].strip())
        ]
        legend = [
            item
            for item in labels
            if item["bbox"][0] >= width * 0.6 and item not in x_labels and item not in y_labels
        ]
        return {
            "labels": labels,
            "direct_values": direct_values,
            "axis_metadata": {"x_labels": x_labels, "y_labels": y_labels},
            "legend_metadata": {"labels": legend},
        }

    def _source(
        self, version_id: UUID
    ) -> tuple[DocumentVersion, Document, Project, StoredFile, Path]:
        row = self.db.execute(
            select(DocumentVersion, Document, Project, StoredFile)
            .join(Document, Document.id == DocumentVersion.document_id)
            .join(Project, Project.id == Document.project_id)
            .join(StoredFile, StoredFile.id == DocumentVersion.source_file_id)
            .where(DocumentVersion.id == version_id)
        ).one_or_none()
        if not row:
            raise AppError(
                code="document_version_not_found", message="文献版本不存在", status_code=404
            )
        version, document, project, stored_file = row
        path = self.storage.path_for_key(stored_file.storage_key)
        if not path.exists():
            raise AppError(code="source_file_missing", message="原始文件不存在", status_code=409)
        return version, document, project, stored_file, path

    def parse(self, version_id: UUID, progress: ProgressCallback) -> dict[str, Any]:
        version, document, project, stored_file, path = self._source(version_id)
        pages = table(self.db, "document_pages")
        self.db.execute(delete(pages).where(pages.c.document_version_id == version_id))
        version.parse_status = "processing"
        self.db.commit()

        extension = (stored_file.extension or path.suffix.lstrip(".")).lower()
        progress(5, "detecting_format")
        if extension == "pdf":
            summary = self._parse_pdf(version, document, project, path, progress)
        elif extension == "docx":
            summary = self._parse_docx(version, path, progress)
        elif extension in {"txt", "md"}:
            summary = self._parse_text(version, path, progress)
        elif extension in {"xlsx", "xls"}:
            if extension == "xls":
                raise AppError(
                    code="legacy_xls_requires_conversion",
                    message="旧版 XLS 请先转换为 XLSX",
                    status_code=422,
                )
            summary = self._parse_xlsx(version, path, progress)
        elif extension == "zip":
            summary = self._parse_zip_manifest(version, path, progress)
        else:
            raise AppError(code="unsupported_parser", message="暂无对应文献解析器", status_code=415)

        version.page_count = summary["pages"]
        version.detected_language = summary.get("language")
        version.parse_status = "partial" if summary.get("requires_ocr") else "completed"
        version.metadata_json = {**(version.metadata_json or {}), "parse_summary": summary}
        self.db.commit()
        progress(100, "completed")
        return summary

    def _insert_page(
        self,
        version_id: UUID,
        page_no: int,
        *,
        text_content: str,
        text_source: str,
        width: float | None = None,
        height: float | None = None,
        rotation: int = 0,
        ocr_confidence: float | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> UUID:
        pages = table(self.db, "document_pages")
        return self.db.execute(
            insert(pages)
            .values(
                document_version_id=version_id,
                page_no=page_no,
                width=width,
                height=height,
                rotation=rotation,
                text_content=text_content,
                text_source=text_source,
                ocr_confidence=ocr_confidence,
                metadata=metadata or {},
            )
            .returning(pages.c.id)
        ).scalar_one()

    def _insert_blocks(
        self,
        version_id: UUID,
        page_id: UUID,
        blocks_data: list[dict[str, Any]],
    ) -> int:
        blocks = table(self.db, "document_blocks")
        rows = []
        for sequence_no, block in enumerate(blocks_data):
            text_value = str(block.get("text") or "").strip()
            if not text_value:
                continue
            rows.append(
                {
                    "document_version_id": version_id,
                    "page_id": page_id,
                    "block_type": block.get("type", "paragraph"),
                    "sequence_no": sequence_no,
                    "section_path": block.get("section_path", []),
                    "content_text": text_value,
                    "bbox": self._bbox(block.get("bbox")),
                    "style": block.get("style", {}),
                    "parser_payload": block.get("payload", {}),
                    "confidence": block.get("confidence"),
                }
            )
        if rows:
            self.db.execute(insert(blocks), rows)
        return len(rows)

    def _parse_pdf(
        self,
        version: DocumentVersion,
        document: Document,
        project: Project,
        path: Path,
        progress: ProgressCallback,
    ) -> dict[str, Any]:
        tables_table = table(self.db, "document_tables")
        cells_table = table(self.db, "document_table_cells")
        figures_table = table(self.db, "document_figures")
        pdf = fitz.open(path)
        metadata = pdf.metadata or {}
        if metadata.get("title") and not document.title:
            document.title = metadata["title"]
        if metadata.get("author") and not document.authors:
            document.authors = [
                item.strip() for item in re.split(r"[;,，；]", metadata["author"]) if item.strip()
            ]

        block_count = table_count = cell_count = figure_count = 0
        requires_ocr = False
        ocr_pages: dict[int, OCRPage] = {}
        ocr_error: str | None = None
        if self.storage.settings.ocr_enabled:
            with tempfile.TemporaryDirectory(prefix="research-pdf-ocr-") as temporary:
                images: list[OCRImage] = []
                for index, page in enumerate(pdf):
                    if page.get_text("text", sort=True).strip():
                        continue
                    page_no = index + 1
                    pixmap = page.get_pixmap(dpi=self.storage.settings.ocr_dpi, alpha=False)
                    image_path = Path(temporary) / f"page-{page_no}.png"
                    pixmap.save(image_path)
                    images.append(
                        OCRImage(
                            page_no=page_no,
                            path=image_path,
                            pixel_width=pixmap.width,
                            pixel_height=pixmap.height,
                            page_width=page.rect.width,
                            page_height=page.rect.height,
                        )
                    )
                if images:
                    progress(8, "running_paddle_ocr")
                    try:
                        ocr_pages = PaddleOCRService(self.storage.settings).recognize(images)
                    except AppError as exc:
                        ocr_error = exc.code

        seen_xrefs: set[int] = set()
        figure_ocr_images: list[OCRImage] = []
        figure_ocr_map: dict[int, UUID] = {}
        figure_ocr_index = 0
        for index, page in enumerate(pdf):
            page_no = index + 1
            raw_blocks = page.get_text("blocks", sort=True)
            text = page.get_text("text", sort=True).strip()
            text_source = "embedded"
            ocr_confidence = None
            if not text:
                paddle_page = ocr_pages.get(page_no)
                if paddle_page and paddle_page.text:
                    text = paddle_page.text
                    raw_blocks = [
                        [*line.bbox, line.text, line_index]
                        for line_index, line in enumerate(paddle_page.lines, start=1)
                    ]
                    text_source = "ocr_paddle"
                    ocr_confidence = paddle_page.confidence
                else:
                    try:
                        text_page = page.get_textpage_ocr(
                            language="chi_sim+eng",
                            dpi=self.storage.settings.ocr_dpi,
                            full=True,
                        )
                        text = page.get_text("text", textpage=text_page, sort=True).strip()
                        raw_blocks = page.get_text("blocks", textpage=text_page, sort=True)
                        text_source = "ocr_tesseract"
                    except Exception:
                        requires_ocr = True
                        text_source = "none"

            page_id = self._insert_page(
                version.id,
                page_no,
                text_content=text,
                text_source=text_source,
                width=page.rect.width,
                height=page.rect.height,
                rotation=page.rotation,
                ocr_confidence=ocr_confidence,
                metadata={
                    "requires_ocr": text_source == "none",
                    "ocr_engine": "PaddleOCR" if text_source == "ocr_paddle" else None,
                },
            )
            parsed_blocks = [
                {
                    "text": item[4],
                    "bbox": item[:4],
                    "type": "paragraph",
                    "payload": {"block_no": item[5]},
                    "confidence": ocr_confidence if text_source == "ocr_paddle" else None,
                }
                for item in raw_blocks
                if len(item) >= 6 and str(item[4]).strip()
            ]
            block_count += self._insert_blocks(version.id, page_id, parsed_blocks)

            try:
                finder = page.find_tables()
                for table_index, found in enumerate(finder.tables, start=1):
                    extracted = found.extract()
                    structure = self._table_structure(found, extracted)
                    title = self._nearby_text(raw_blocks, found.bbox, above=True)
                    footnote = self._nearby_text(raw_blocks, found.bbox, above=False)
                    footnotes = (
                        [footnote]
                        if footnote and re.match(r"^(?:注|Note|\*|[a-z]\))", footnote, re.I)
                        else []
                    )
                    structure["footnotes"] = footnotes
                    table_id = self.db.execute(
                        insert(tables_table)
                        .values(
                            document_version_id=version.id,
                            page_id=page_id,
                            table_no=f"p{page_no}-t{table_index}",
                            title=title,
                            caption=title,
                            row_count=len(extracted),
                            column_count=max((len(row) for row in extracted), default=0),
                            bbox=self._bbox(found.bbox),
                            structured_data=structure,
                            confidence=0.9,
                        )
                        .returning(tables_table.c.id)
                    ).scalar_one()
                    table_count += 1
                    cell_rows = []
                    for row_index, row in enumerate(extracted):
                        for column_index, value in enumerate(row):
                            geometry = None
                            if row_index < len(found.rows) and column_index < len(
                                found.rows[row_index].cells
                            ):
                                geometry = found.rows[row_index].cells[column_index]
                            merged_cell: dict[str, Any] = next(
                                (
                                    item
                                    for item in structure["merged_cells"]
                                    if item["row"] == row_index and item["column"] == column_index
                                ),
                                {},
                            )
                            text_value = value.strip() if isinstance(value, str) else value
                            significance = None
                            if isinstance(text_value, str):
                                marker = re.search(r"(?<=\d)([a-zA-Z]{1,3})$", text_value)
                                significance = marker.group(1) if marker else None
                            cell_rows.append(
                                {
                                    "table_id": table_id,
                                    "row_index": row_index,
                                    "column_index": column_index,
                                    "row_span": merged_cell.get("row_span", 1),
                                    "column_span": merged_cell.get("column_span", 1),
                                    "cell_role": "header"
                                    if row_index < structure["header_rows"]
                                    else "data",
                                    "raw_text": value,
                                    "normalized_text": text_value,
                                    "bbox": self._bbox(geometry),
                                    "style": {
                                        "header_path": structure["header_paths"][column_index]
                                        if column_index < len(structure["header_paths"])
                                        else [],
                                        "unit": structure["units"].get(str(column_index)),
                                        "significance_marker": significance,
                                    },
                                    "confidence": 0.9,
                                }
                            )
                    if cell_rows:
                        self.db.execute(insert(cells_table), cell_rows)
                        cell_count += len(cell_rows)
            except Exception:
                pass

            for image_index, image in enumerate(page.get_images(full=True), start=1):
                xref = int(image[0])
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)
                image_file_id = None
                saved = None
                extracted_image = None
                figure_bbox = None
                try:
                    extracted_image = pdf.extract_image(xref)
                    content = extracted_image["image"]
                    extension = extracted_image.get("ext", "png")
                    saved = self.storage.save_bytes(
                        project.id,
                        category="figures",
                        extension=extension,
                        content=content,
                        media_type=mimetypes.guess_type(f"figure.{extension}")[0],
                    )
                    stored = StoredFile(
                        organization_id=project.organization_id,
                        project_id=project.id,
                        storage_provider="local",
                        storage_key=saved.storage_key,
                        original_name=f"page-{page_no}-image-{image_index}.{extension}",
                        safe_name=saved.safe_name,
                        extension=extension,
                        media_type=saved.media_type,
                        byte_size=saved.byte_size,
                        sha256=saved.sha256,
                        purpose="figure",
                        security_status="generated",
                        metadata_json={"source_xref": xref},
                    )
                    self.db.add(stored)
                    self.db.flush()
                    image_file_id = stored.id
                    rects = page.get_image_rects(xref)
                    figure_bbox = self._bbox(rects[0]) if rects else None
                except Exception:
                    image_file_id = None
                caption = (
                    self._nearby_text(raw_blocks, figure_bbox, above=False, distance=110)
                    if figure_bbox
                    else None
                )
                figure_id = self.db.execute(
                    insert(figures_table)
                    .values(
                        document_version_id=version.id,
                        page_id=page_id,
                        figure_no=f"p{page_no}-f{image_index}",
                        title=caption,
                        caption=caption,
                        figure_type="embedded_image",
                        bbox=figure_bbox,
                        image_file_id=image_file_id,
                        axis_metadata={},
                        legend_metadata={},
                        extracted_labels=[],
                        confidence=0.7,
                    )
                    .returning(figures_table.c.id)
                ).scalar_one()
                if saved and extracted_image:
                    figure_ocr_index += 1
                    figure_ocr_images.append(
                        OCRImage(
                            page_no=figure_ocr_index,
                            path=saved.path,
                            pixel_width=int(extracted_image.get("width") or 1),
                            pixel_height=int(extracted_image.get("height") or 1),
                            page_width=float(extracted_image.get("width") or 1),
                            page_height=float(extracted_image.get("height") or 1),
                        )
                    )
                    figure_ocr_map[figure_ocr_index] = figure_id
                figure_count += 1

            progress(5 + 90 * page_no / max(len(pdf), 1), f"parsing_page_{page_no}")
            self.db.commit()
        if figure_ocr_images and self.storage.settings.ocr_enabled:
            try:
                progress(96, "analyzing_figure_labels")
                figure_pages = PaddleOCRService(self.storage.settings).recognize(figure_ocr_images)
                for ocr_index, ocr_page in figure_pages.items():
                    figure_id = figure_ocr_map.get(ocr_index)
                    if not figure_id:
                        continue
                    image = next(item for item in figure_ocr_images if item.page_no == ocr_index)
                    metadata = self._figure_metadata(
                        ocr_page.lines, image.page_width, image.page_height
                    )
                    self.db.execute(
                        update(figures_table)
                        .where(figures_table.c.id == figure_id)
                        .values(
                            axis_metadata=metadata["axis_metadata"],
                            legend_metadata=metadata["legend_metadata"],
                            extracted_labels={
                                "labels": metadata["labels"],
                                "direct_values": metadata["direct_values"],
                                "values_are_estimated": False,
                            },
                            semantic_summary=" | ".join(
                                item["text"] for item in metadata["labels"]
                            ),
                            confidence=ocr_page.confidence,
                        )
                    )
                self.db.commit()
            except AppError:
                pass
        page_total = len(pdf)
        pdf.close()
        return {
            "pages": page_total,
            "blocks": block_count,
            "tables": table_count,
            "cells": cell_count,
            "figures": figure_count,
            "requires_ocr": requires_ocr,
            "ocr_engine": "PaddleOCR" if ocr_pages else None,
            "ocr_page_count": len(ocr_pages),
            "ocr_error": ocr_error,
            "language": version.detected_language,
        }

    def _parse_docx(
        self, version: DocumentVersion, path: Path, progress: ProgressCallback
    ) -> dict[str, Any]:
        document = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        page_id = self._insert_page(version.id, 1, text_content=text, text_source="embedded")
        blocks = [{"text": value, "type": "paragraph"} for value in paragraphs]
        block_count = self._insert_blocks(version.id, page_id, blocks)
        tables_table = table(self.db, "document_tables")
        cells_table = table(self.db, "document_table_cells")
        cell_count = 0
        for index, docx_table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in docx_table.rows]
            table_id = self.db.execute(
                insert(tables_table)
                .values(
                    document_version_id=version.id,
                    page_id=page_id,
                    table_no=f"docx-t{index}",
                    row_count=len(rows),
                    column_count=max((len(row) for row in rows), default=0),
                    structured_data={"rows": rows},
                    confidence=0.95,
                )
                .returning(tables_table.c.id)
            ).scalar_one()
            values = [
                {
                    "table_id": table_id,
                    "row_index": row_no,
                    "column_index": col_no,
                    "cell_role": "header" if row_no == 0 else "data",
                    "raw_text": value,
                    "normalized_text": value,
                    "style": {},
                    "confidence": 0.95,
                }
                for row_no, row in enumerate(rows)
                for col_no, value in enumerate(row)
            ]
            if values:
                self.db.execute(insert(cells_table), values)
                cell_count += len(values)
        self.db.commit()
        progress(95, "parsed_docx")
        return {
            "pages": 1,
            "blocks": block_count,
            "tables": len(document.tables),
            "cells": cell_count,
            "figures": 0,
            "requires_ocr": False,
        }

    def _parse_text(
        self, version: DocumentVersion, path: Path, progress: ProgressCallback
    ) -> dict[str, Any]:
        raw = path.read_bytes()
        text = raw.decode("utf-8-sig", errors="replace")
        paragraphs = [
            item.strip() for item in re.split(r"\n\s*\n|(?<=。)\s*", text) if item.strip()
        ]
        page_id = self._insert_page(version.id, 1, text_content=text, text_source="embedded")
        block_count = self._insert_blocks(version.id, page_id, [{"text": p} for p in paragraphs])
        self.db.commit()
        progress(95, "parsed_text")
        return {
            "pages": 1,
            "blocks": block_count,
            "tables": 0,
            "cells": 0,
            "figures": 0,
            "requires_ocr": False,
        }

    def _parse_xlsx(
        self, version: DocumentVersion, path: Path, progress: ProgressCallback
    ) -> dict[str, Any]:
        workbook = load_workbook(path, read_only=True, data_only=True)
        tables_table = table(self.db, "document_tables")
        cells_table = table(self.db, "document_table_cells")
        total_cells = 0
        for page_no, sheet in enumerate(workbook.worksheets, start=1):
            rows = [
                ["" if value is None else str(value) for value in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            text = "\n".join("\t".join(row) for row in rows)
            page_id = self._insert_page(
                version.id,
                page_no,
                text_content=text,
                text_source="spreadsheet",
                metadata={"sheet_name": sheet.title},
            )
            self._insert_blocks(version.id, page_id, [{"text": text, "type": "table_text"}])
            table_id = self.db.execute(
                insert(tables_table)
                .values(
                    document_version_id=version.id,
                    page_id=page_id,
                    table_no=sheet.title,
                    title=sheet.title,
                    row_count=len(rows),
                    column_count=max((len(row) for row in rows), default=0),
                    structured_data={"rows": rows},
                    confidence=1,
                )
                .returning(tables_table.c.id)
            ).scalar_one()
            values = [
                {
                    "table_id": table_id,
                    "row_index": row_no,
                    "column_index": col_no,
                    "cell_role": "header" if row_no == 0 else "data",
                    "raw_text": value,
                    "normalized_text": value,
                    "style": {},
                    "confidence": 1,
                }
                for row_no, row in enumerate(rows)
                for col_no, value in enumerate(row)
            ]
            if values:
                self.db.execute(insert(cells_table), values)
                total_cells += len(values)
            progress(
                5 + 90 * page_no / max(len(workbook.worksheets), 1), f"parsing_sheet_{page_no}"
            )
            self.db.commit()
        return {
            "pages": len(workbook.worksheets),
            "blocks": len(workbook.worksheets),
            "tables": len(workbook.worksheets),
            "cells": total_cells,
            "figures": 0,
            "requires_ocr": False,
        }

    def _parse_zip_manifest(
        self, version: DocumentVersion, path: Path, progress: ProgressCallback
    ) -> dict[str, Any]:
        import zipfile

        with zipfile.ZipFile(path) as archive:
            names = [name for name in archive.namelist() if not name.endswith("/")]
        text = "\n".join(names)
        page_id = self._insert_page(
            version.id,
            1,
            text_content=text,
            text_source="archive_manifest",
            metadata={"archive_entries": names},
        )
        count = self._insert_blocks(
            version.id, page_id, [{"text": name, "type": "archive_entry"} for name in names]
        )
        self.db.commit()
        progress(95, "parsed_archive_manifest")
        return {
            "pages": 1,
            "blocks": count,
            "tables": 0,
            "cells": 0,
            "figures": 0,
            "requires_ocr": False,
            "archive_only": True,
        }
